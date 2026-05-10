from google import genai
from google.genai import types
import os
import json
import time
from dotenv import load_dotenv
from docx import Document
import tempfile

load_dotenv()
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY', '')

def get_best_model():
    client = genai.Client(api_key=GOOGLE_API_KEY)
    '''try:
        for m in client.models.list():
            if m.supported_actions and 'generateContent' in m.supported_actions and 'flash' in m.name:
                return m.name
    except: pass'''
    return 'gemini-3.1-flash-lite'

def convert_docx_to_txt(docx_path):
    """Convert DOCX file to TXT by extracting all text"""
    try:
        doc = Document(docx_path)
        txt_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                txt_content.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        txt_content.append(cell.text)
        
        # Create temp TXT file
        txt_path = tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8')
        txt_path.write('\n'.join(txt_content))
        txt_path.close()
        print(f"✅ DOCX converted to TXT: {txt_path.name} ({len(txt_content)} lines extracted)")
        return txt_path.name
    except Exception as e:
        print(f"❌ DOCX Conversion Error: {e}")
        return None

def get_mime_type(file_path):
    """Get the MIME type based on file extension"""
    ext = os.path.splitext(file_path)[1].lower()
    mime_types = {
        '.pdf': 'application/pdf',
        '.txt': 'text/plain',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
    }
    return mime_types.get(ext, 'application/octet-stream')

def analyze_document_with_gemini(file_path, available_departments=None):
    print(f"--- 🧠 Starting AI Analysis for: {file_path} ---")
    client = genai.Client(api_key=GOOGLE_API_KEY)
    
    # Convert DOCX to TXT if needed
    analysis_file_path = file_path
    temp_txt_path = None
    if file_path.lower().endswith('.docx'):
        print(f"📄 Converting DOCX to TXT...")
        temp_txt_path = convert_docx_to_txt(file_path)
        if temp_txt_path:
            analysis_file_path = temp_txt_path
        else:
            return []
    
    try:
        mime_type = get_mime_type(analysis_file_path)
        print(f"📄 File MIME type: {mime_type}")
        uploaded_file = client.files.upload(file=analysis_file_path, config={'display_name': 'Minutes', 'mime_type': mime_type})
        while uploaded_file.state.name == "PROCESSING":
            time.sleep(1)
            uploaded_file = client.files.get(name=uploaded_file.name)
    except Exception as e:
        print(f"❌ Upload Error: {e}")
        # Clean up temp file
        if temp_txt_path and os.path.exists(temp_txt_path):
            os.remove(temp_txt_path)
        return []

    safety = [
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_HARASSMENT, threshold=types.HarmBlockThreshold.BLOCK_NONE),
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_HATE_SPEECH, threshold=types.HarmBlockThreshold.BLOCK_NONE),
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT, threshold=types.HarmBlockThreshold.BLOCK_NONE),
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT, threshold=types.HarmBlockThreshold.BLOCK_NONE),
    ]
    
    available_departments = available_departments or []
    dept_list_text = '\n'.join(available_departments) if available_departments else '(No department list provided)'

    prompt = f"""
    You are an AI specialized in analyzing Malayalam minutes of meetings (PDF, DOCX, etc.). 

    TASK: Extract every actionable issue from this document.

    AVAILABLE DEPARTMENTS (Designation, Department Name):
    {dept_list_text}

    EXTRACTION RULES:
    1. Extract ONLY issues that have (നടപടി at the END, which marks it as an actionable item.
    2. The pattern is: (നടപടി followed by optional spaces, then optional punctuation (: or –), then optional spaces, then department/designation info.
    3. Examples: (നടപടി:, (നടപടി –, (നടപടി  :, (നടപടി – etc. ALL are valid markers.
    4. CRITICAL: Ignore നടപടി if it appears in the MIDDLE of the issue text - only extract when preceded by ( and followed by department info.
    5. For each (നടപടി marker, identify ALL mentioned stakeholders (officers/departments) that follow it.
    6. For EACH stakeholder, extract BOTH the "designation" and the "department" name as written in the text.
    7. CRITICAL INSTRUCTION: You MUST ONLY return a department if it is EXPLICITLY stated in the text. 
    8. ZERO GUESSING ALLOWED: Do NOT add related departments. Do NOT infer a department just because of the context.
    9. Map each extracted pair to the AVAILABLE DEPARTMENTS list.
    10. Do not invent any dates. If no deadline is explicitly given, set "deadline" to "".
    11. Maintain the original Malayalam text for issue_description and location fields.
    12. Summarize each issue in ONE line (max 20 words) for the "issue" field.
    13. Assign priority: High/Medium/Low.

    CRITICAL NEGATIVE EXAMPLES (DO NOT DO THIS):
    - If the (നടപടി section says "വാട്ടർ അതോറിറ്റി" (Water Authority), ONLY return "Water Authority".
    - If it says "എക്സിക്യുട്ടീവ് എൻജിനീയർ, പൊതുമരാമത്ത് (കെട്ടിടം)", extract Designation: "എക്സിക്യുട്ടീവ് എൻജിനീയർ", Department: "പൊതുമരാമത്ത് (കെട്ടിടം)".
    - If നടപടി appears in the MIDDLE of text (not after () as an action marker), IGNORE that issue completely.
    OUTPUT FORMAT: Return ONLY a JSON array of objects. Each object must have:

    {{
    "issue_no": string,
    "departments": [
        {{
            "designation": "english version of Malayalam designation extracted from text",
            "department": "english version of Malayalam department name extracted from text"
        }}
    ],
    "issue": one-line Malayalam summary (max 20 words),
    "issue_description": full Malayalam issue description exactly as in document,
    "location": Malayalam place name if mentioned, else "",
    "priority": High/Medium/Low,
    "deadline": DD-MM-YYYY or ""
    }}

    Do NOT include any text outside the JSON array. Do NOT add explanations or code blocks.
    """

    try:
        response = client.models.generate_content(
            model=get_best_model(),
            contents=[uploaded_file, prompt],
            config=types.GenerateContentConfig(safety_settings=safety)
        )
        text = response.text.replace("```json", "").replace("```", "").strip()
        result = json.loads(text)
        # Clean up temp file
        if temp_txt_path and os.path.exists(temp_txt_path):
            os.remove(temp_txt_path)
        return result
    except Exception as e:
        print(f"❌ Generation Error: {e}")
        # Clean up temp file
        if temp_txt_path and os.path.exists(temp_txt_path):
            os.remove(temp_txt_path)
        return []


def match_issues_with_gemini(new_issues, existing_issues):
    """
    Use Gemini to semantically match new issues against existing unresolved issues.
    Works with Malayalam text — compares meaning, not exact words.
    
    new_issues: list of dicts with keys: index, issue, issue_description
    existing_issues: list of dicts with keys: id, issue, issue_description, minutes_title
    
    Returns: list of dicts { new_index: int, existing_id: int, confidence: str }
    """
    if not new_issues or not existing_issues:
        return []

    print(f"--- 🔗 Starting Issue Matching: {len(new_issues)} new vs {len(existing_issues)} existing ---")
    client = genai.Client(api_key=GOOGLE_API_KEY)

    safety = [
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_HARASSMENT, threshold=types.HarmBlockThreshold.BLOCK_NONE),
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_HATE_SPEECH, threshold=types.HarmBlockThreshold.BLOCK_NONE),
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT, threshold=types.HarmBlockThreshold.BLOCK_NONE),
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT, threshold=types.HarmBlockThreshold.BLOCK_NONE),
    ]

    prompt = f"""
You are an expert at matching government meeting issues written in Malayalam.

TASK: Compare NEW issues (extracted from the latest meeting minutes) against EXISTING issues (from previous meetings). Find which new issues are follow-ups or continuations of existing issues.

NEW ISSUES:
{json.dumps(new_issues, ensure_ascii=False, indent=2)}

EXISTING ISSUES:
{json.dumps(existing_issues, ensure_ascii=False, indent=2)}

MATCHING RULES:
1. Match based on SEMANTIC MEANING, not exact text. The same issue may be described differently across meetings.
2. Consider: same road/location, same infrastructure problem, same department concern, same complaint.
3. Only match if you are reasonably confident they refer to the SAME real-world issue.
4. A new issue can match AT MOST one existing issue.
5. Not every new issue will have a match — only return genuine matches.
6. confidence must be "high" (clearly same issue) or "medium" (likely same issue).

Return a JSON ARRAY of match objects:
- new_index: integer (the "index" field from the new issue)
- existing_id: integer (the "id" field from the matched existing issue)
- confidence: "high" or "medium"

If no matches found, return an empty array [].
Output ONLY PURE JSON.
"""

    try:
        response = client.models.generate_content(
            model=get_best_model(),
            contents=prompt,
            config=types.GenerateContentConfig(safety_settings=safety)
        )
        text = response.text.replace("```json", "").replace("```", "").strip()
        matches = json.loads(text)
        print(f"✅ Gemini returned {len(matches)} matches")
        return matches if isinstance(matches, list) else []
    except Exception as e:
        print(f"❌ Matching Error: {e}")
        return []